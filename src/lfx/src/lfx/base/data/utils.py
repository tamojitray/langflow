import contextlib
import tempfile
import unicodedata
from collections.abc import Callable
from concurrent import futures
from io import BytesIO
from pathlib import Path

import chardet
import orjson
import yaml
from defusedxml import ElementTree
from pypdf import PdfReader

from lfx.base.data.storage_utils import read_file_bytes
from lfx.schema.data import Data
from lfx.services.deps import get_settings_service
from lfx.utils.async_helpers import run_until_complete

# Types of files that can be read simply by file.read()
# and have 100% to be completely readable
TEXT_FILE_TYPES = [
    "csv",
    "doc",
    "docx",
    "json",
    "pdf",
    "rtf",
    "txt",
    "md",
    "mdx",
    "yaml",
    "yml",
    "xml",
    "html",
    "htm",
    "py",
    "sh",
    "sql",
    "js",
    "ts",
    "tsx",
]

IMG_FILE_TYPES = ["jpg", "jpeg", "png", "bmp", "image"]


def parse_structured_text(text: str, file_path: str) -> str | dict | list:
    """Parse structured text formats (JSON, YAML, XML) and normalize text.

    Args:
        text: The text content to parse
        file_path: The file path (used to determine format)

    Returns:
        Parsed content (dict/list for JSON, dict for YAML, str for XML)
    """
    if file_path.endswith(".json"):
        loaded_json = orjson.loads(text)
        if isinstance(loaded_json, dict):
            loaded_json = {k: normalize_text(v) if isinstance(v, str) else v for k, v in loaded_json.items()}
        elif isinstance(loaded_json, list):
            loaded_json = [normalize_text(item) if isinstance(item, str) else item for item in loaded_json]
        return orjson.dumps(loaded_json).decode("utf-8")

    if file_path.endswith((".yaml", ".yml")):
        return yaml.safe_load(text)

    if file_path.endswith(".xml"):
        xml_element = ElementTree.fromstring(text)
        return ElementTree.tostring(xml_element, encoding="unicode")

    return text


def normalize_text(text):
    return unicodedata.normalize("NFKD", text)


def is_hidden(path: Path) -> bool:
    return path.name.startswith(".")


def format_directory_path(path: str) -> str:
    """Format a directory path to ensure it's properly escaped and valid.

    Args:
    path (str): The input path string.

    Returns:
    str: A properly formatted path string.
    """
    return path.replace("\n", "\\n")


# Ignoring FBT001 because the DirectoryComponent in 1.0.19
# calls this function without keyword arguments
def retrieve_file_paths(
    path: str,
    load_hidden: bool,  # noqa: FBT001
    recursive: bool,  # noqa: FBT001
    depth: int,
    types: list[str] = TEXT_FILE_TYPES,
) -> list[str]:
    path = format_directory_path(path)
    path_obj = Path(path)
    if not path_obj.exists() or not path_obj.is_dir():
        msg = f"Path {path} must exist and be a directory."
        raise ValueError(msg)

    def match_types(p: Path) -> bool:
        return any(p.suffix == f".{t}" for t in types) if types else True

    def is_not_hidden(p: Path) -> bool:
        return not is_hidden(p) or load_hidden

    def walk_level(directory: Path, max_depth: int):
        directory = directory.resolve()
        prefix_length = len(directory.parts)
        for p in directory.rglob("*" if recursive else "[!.]*"):
            if len(p.parts) - prefix_length <= max_depth:
                yield p

    glob = "**/*" if recursive else "*"
    paths = walk_level(path_obj, depth) if depth else path_obj.glob(glob)
    return [str(p) for p in paths if p.is_file() and match_types(p) and is_not_hidden(p)]


def partition_file_to_data(file_path: str, *, silent_errors: bool) -> Data | None:
    # Use the partition function to load the file
    from unstructured.partition.auto import partition

    try:
        elements = partition(file_path)
    except Exception as e:
        if not silent_errors:
            msg = f"Error loading file {file_path}: {e}"
            raise ValueError(msg) from e
        return None

    # Create a Data
    text = "\n\n".join([str(el) for el in elements])
    metadata = elements.metadata if hasattr(elements, "metadata") else {}
    metadata["file_path"] = file_path
    return Data(text=text, data=metadata)


def read_text_file(file_path: str) -> str:
    """Read a text file with automatic encoding detection.

    Args:
        file_path: Path to the file (local path only, not storage service path)

    Returns:
        str: The file content as text
    """
    file_path_ = Path(file_path)
    raw_data = file_path_.read_bytes()
    result = chardet.detect(raw_data)
    encoding = result["encoding"]

    if encoding in {"Windows-1252", "Windows-1254", "MacRoman"}:
        encoding = "utf-8"

    return file_path_.read_text(encoding=encoding)


async def read_text_file_async(file_path: str) -> str:
    """Read a text file with automatic encoding detection (async, storage-aware).

    Args:
        file_path: Path to the file (S3 key format "flow_id/filename" or local path)

    Returns:
        str: The file content as text
    """
    from .storage_utils import read_file_bytes

    # Use storage-aware read to get bytes
    raw_data = await read_file_bytes(file_path)

    # Auto-detect encoding
    result = chardet.detect(raw_data)
    encoding = result.get("encoding")

    # If encoding detection fails (e.g., binary file), default to utf-8
    if not encoding or encoding in {"Windows-1252", "Windows-1254", "MacRoman"}:
        encoding = "utf-8"

    return raw_data.decode(encoding, errors="replace")


def _extract_docx_full_content(doc) -> str:
    """Extract ALL content from a python-docx Document in reading order.

    This walks the document body XML to yield paragraphs, tables, headers,
    and footers in the order they appear.  Tables are rendered as
    pipe-delimited Markdown-style tables so that their data is fully
    captured in the text output.
    """
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    parts: list[str] = []

    # --- Helper: render a single table to markdown-style text ---------------
    def _table_to_text(table) -> str:
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows_text.append("| " + " | ".join(cells) + " |")
        if len(rows_text) >= 1:
            # Insert a header separator after the first row
            col_count = len(table.rows[0].cells)
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            rows_text.insert(1, separator)
        return "\n".join(rows_text)

    # --- Walk body elements in document order --------------------------------
    # doc.element.body contains <w:p> (paragraph) and <w:tbl> (table) children.
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para = DocxParagraph(child, doc.element.body)
            text = para.text.strip()
            if text:
                parts.append(text)
        elif tag == "tbl":
            table = DocxTable(child, doc.element.body)
            rendered = _table_to_text(table)
            if rendered:
                parts.append(rendered)

    # --- Extract headers and footers ----------------------------------------
    for section in doc.sections:
        for hf_attr in ("header", "first_page_header", "even_page_header",
                        "footer", "first_page_footer", "even_page_footer"):
            hf = getattr(section, hf_attr, None)
            if hf is None or not hf.is_linked_to_previous:
                continue
            for para in hf.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

    return "\n\n".join(parts)


def read_docx_file(file_path: str) -> str:
    """Read a DOCX file and extract ALL text including tables.

    Uses _extract_docx_full_content to walk the XML body so that tables,
    paragraphs, headers, and footers are all captured.

    Note: python-docx requires a file path, so this only works with local files.
    For storage service files, use read_docx_file_async which downloads to temp.

    Args:
        file_path: Path to the DOCX file (local path only)

    Returns:
        str: Extracted text from the document
    """
    from docx import Document

    doc = Document(file_path)
    return _extract_docx_full_content(doc)


def read_doc_file(file_path: str) -> str:
    """Read a legacy .doc (Word 97-2003) file and extract text.

    Tries multiple strategies:
    1. ``textract`` (if installed) – most reliable on all platforms.
    2. ``antiword`` via subprocess (common on Linux).
    3. Falls back to reading raw bytes and stripping non-text.

    Args:
        file_path: Path to the .doc file (local path only)

    Returns:
        str: Extracted text from the document
    """
    # Strategy 1: textract
    try:
        import textract  # type: ignore[import-untyped]

        raw = textract.process(file_path)
        return raw.decode("utf-8", errors="replace")
    except ImportError:
        pass
    except Exception:
        pass

    # Strategy 2: antiword (commonly available on Linux)
    import subprocess  # noqa: S404

    try:
        result = subprocess.run(  # noqa: S603, S607
            ["antiword", file_path],
            capture_output=True,
            text=True,
            check=True,
        )
        return result.stdout
    except (FileNotFoundError, subprocess.CalledProcessError):
        pass

    # Strategy 3: raw byte extraction (last resort – lossy)
    import re

    raw = Path(file_path).read_bytes()
    # The .doc OLE2 format stores text between certain markers; strip control chars.
    text = raw.decode("utf-8", errors="replace")
    # Remove non-printable characters but keep newlines/tabs
    text = re.sub(r"[^\x20-\x7E\n\r\t]", "", text)
    # Collapse excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    if text.strip():
        return text.strip()

    msg = (
        f"Could not extract text from legacy .doc file '{file_path}'. "
        "Install 'textract' (pip install textract) or 'antiword' for .doc support, "
        "or convert the file to .docx format."
    )
    raise ValueError(msg)


def _read_rtf_file(file_path: str) -> str:
    """Read an RTF file and extract plain text.

    Tries ``striprtf`` first (pip install striprtf), then falls back to a
    basic regex approach that strips RTF control words.

    Args:
        file_path: Path to the .rtf file

    Returns:
        str: Extracted plain text
    """
    import re as _re

    raw = Path(file_path).read_bytes().decode("utf-8", errors="replace")

    # Strategy 1: striprtf library
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore[import-untyped]

        return rtf_to_text(raw)
    except ImportError:
        pass

    # Strategy 2: basic regex stripping of RTF control codes
    # Remove RTF groups like {\fonttbl ...}, {\colortbl ...}, etc.
    text = _re.sub(r"\{\\[a-z]+[^}]*\}", "", raw)
    # Remove RTF control words like \par, \b0, \fs24, etc.
    text = _re.sub(r"\\[a-z]+\d*\s?", " ", text)
    # Remove remaining braces
    text = _re.sub(r"[{}]", "", text)
    # Collapse whitespace
    text = _re.sub(r"[ \t]+", " ", text)
    text = _re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


async def read_docx_file_async(file_path: str) -> str:
    """Read a DOCX file and extract ALL text including tables (async, storage-aware).

    For S3 storage, downloads to temp file (python-docx requires file path).
    For local storage, reads directly.

    Args:
        file_path: Path to the DOCX file (S3 key format "flow_id/filename" or local path)

    Returns:
        str: Extracted text from the document
    """
    from docx import Document

    from .storage_utils import read_file_bytes

    settings = get_settings_service().settings

    if settings.storage_type == "local":
        # Local storage - read directly
        doc = Document(file_path)
        return _extract_docx_full_content(doc)

    # S3 storage - need temp file for python-docx (doesn't support BytesIO)
    content = await read_file_bytes(file_path)

    # Create temp file with .docx extension
    # Extract filename from path for suffix
    suffix = Path(file_path.split("/")[-1]).suffix
    with tempfile.NamedTemporaryFile(mode="wb", suffix=suffix, delete=False) as tmp_file:
        tmp_file.write(content)
        temp_path = tmp_file.name

    try:
        doc = Document(temp_path)
        return _extract_docx_full_content(doc)
    finally:
        with contextlib.suppress(Exception):
            Path(temp_path).unlink()


def extract_text_from_bytes(file_name: str, file_content: bytes) -> str:
    """Extract text from binary file content based on file extension.

    Supports PDF (via pypdf), DOCX (via python-docx with full table extraction),
    and plain text files.

    Raises:
        ValueError: If the file content is corrupted or cannot be parsed.
    """
    lower_name = file_name.lower()
    if lower_name.endswith(".pdf"):
        try:
            with BytesIO(file_content) as f, PdfReader(f) as reader:
                return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            msg = f"Failed to parse PDF file '{file_name}': {e}"
            raise ValueError(msg) from e
    if lower_name.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(BytesIO(file_content))
            return _extract_docx_full_content(doc)
        except Exception as e:
            msg = f"Failed to parse DOCX file '{file_name}': {e}"
            raise ValueError(msg) from e
    return file_content.decode("utf-8", errors="ignore")


def parse_pdf_to_text(file_path: str) -> str:
    with Path(file_path).open("rb") as f, PdfReader(f) as reader:
        return "\n\n".join([page.extract_text() for page in reader.pages])


async def parse_pdf_to_text_async(file_path: str) -> str:
    """Parse a PDF file to extract text (async, storage-aware).

    Uses storage-aware file reading to support both local and S3 storage.

    Args:
        file_path: Path to the PDF file (S3 key format "flow_id/filename" or local path)

    Returns:
        str: Extracted text from all pages
    """
    content = await read_file_bytes(file_path)
    with BytesIO(content) as f, PdfReader(f) as reader:
        return "\n\n".join([page.extract_text() for page in reader.pages])


def parse_text_file_to_data(file_path: str, *, silent_errors: bool) -> Data | None:
    """Parse a text file to Data (sync version).

    For S3 storage, this will use async operations to fetch the file.
    For local storage, reads directly from filesystem.
    """
    settings = get_settings_service().settings

    # If using S3 storage, we need to use async operations
    if settings.storage_type == "s3":
        # Run the async version safely (handles existing event loops)
        return run_until_complete(parse_text_file_to_data_async(file_path, silent_errors=silent_errors))

    try:
        if file_path.endswith(".pdf"):
            text = parse_pdf_to_text(file_path)
        elif file_path.endswith(".docx"):
            text = read_docx_file(file_path)
        elif file_path.endswith(".doc"):
            text = read_doc_file(file_path)
        elif file_path.endswith(".rtf"):
            text = _read_rtf_file(file_path)
        else:
            text = read_text_file(file_path)

        text = parse_structured_text(text, file_path)
    except Exception as e:
        if not silent_errors:
            msg = f"Error loading file {file_path}: {e}"
            raise ValueError(msg) from e
        return None

    return Data(data={"file_path": file_path, "text": text})


async def parse_text_file_to_data_async(file_path: str, *, silent_errors: bool) -> Data | None:
    """Parse a text file to Data (async version, supports storage service).

    This version properly handles storage service files:
    - For text/JSON/YAML/XML: reads bytes directly (no temp file)
    - For PDF: reads bytes directly via BytesIO (no temp file)
    - For DOCX: downloads to temp file (python-docx requires file path)
    """
    try:
        if file_path.endswith(".pdf"):
            text = await parse_pdf_to_text_async(file_path)
        elif file_path.endswith(".docx"):
            text = await read_docx_file_async(file_path)
        elif file_path.endswith(".doc"):
            # .doc files always need a local path; download from storage if needed
            text = read_doc_file(file_path)
        elif file_path.endswith(".rtf"):
            text = _read_rtf_file(file_path)
        else:
            # Text files - read directly, no temp file needed
            text = await read_text_file_async(file_path)

        # Parse structured formats (JSON, YAML, XML)
        text = parse_structured_text(text, file_path)

        return Data(data={"file_path": file_path, "text": text})

    except Exception as e:
        if not silent_errors:
            msg = f"Error loading file {file_path}: {e}"
            raise ValueError(msg) from e
        return None


# ! Removing unstructured dependency until
# ! 3.12 is supported
# def get_elements(
#     file_paths: List[str],
#     silent_errors: bool,
#     max_concurrency: int,
#     use_multithreading: bool,
# ) -> List[Optional[Data]]:
#     if use_multithreading:
#         data = parallel_load_data(file_paths, silent_errors, max_concurrency)
#     else:
#         data = [partition_file_to_data(file_path, silent_errors) for file_path in file_paths]
#     data = list(filter(None, data))
#     return data


def parallel_load_data(
    file_paths: list[str],
    *,
    silent_errors: bool,
    max_concurrency: int,
    load_function: Callable = parse_text_file_to_data,
) -> list[Data | None]:
    with futures.ThreadPoolExecutor(max_workers=max_concurrency) as executor:
        loaded_files = executor.map(
            lambda file_path: load_function(file_path, silent_errors=silent_errors),
            file_paths,
        )
    # loaded_files is an iterator, so we need to convert it to a list
    return list(loaded_files)
